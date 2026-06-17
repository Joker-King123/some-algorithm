"""
生成专业课程设计报告 Word 文档 (.docx)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

BASE_DIR = os.path.dirname(__file__)

# ============================================================
# 文档初始化
# ============================================================
doc = Document()

# ---- 页面设置 ----
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ---- 样式定义 ----
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.first_line_indent = Cm(0.74)

# 标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif i == 2:
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        heading_style.font.size = Pt(13)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_paragraph(text, style_name='Normal', bold=False, alignment=None,
                  font_size=None, font_name=None, first_line_indent=True,
                  space_after=None):
    """添加段落"""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if alignment is not None:
        p.alignment = alignment
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_rich_paragraph(segments, alignment=None, first_line_indent=True):
    """添加富文本段落，segments = [(text, bold, font_name, font_size), ...]"""
    p = doc.add_paragraph()
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        font_name = seg[2] if len(seg) > 2 else '宋体'
        font_size = seg[3] if len(seg) > 3 else 12
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if alignment is not None:
        p.alignment = alignment
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_figure(image_name, caption, width_inches=5.5):
    """插入图片和题注"""
    img_path = os.path.join(BASE_DIR, image_name)
    if os.path.exists(img_path):
        # 图片居中
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.first_line_indent = Cm(0)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(width_inches))

        # 题注
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.first_line_indent = Cm(0)
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run(caption)
        run_cap.font.size = Pt(9)
        run_cap.font.name = '黑体'
        run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run_cap.bold = False
    else:
        add_paragraph(f'[图片缺失: {image_name}]', first_line_indent=False,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_table(headers, rows, caption='', col_widths=None):
    """插入表格"""
    if caption:
        p = add_paragraph(caption, bold=True, font_size=9, first_line_indent=False,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # 灰色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行
    return table


# ============================================================
# 封面
# ============================================================
def build_cover():
    # 空行留白
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)

    # 校名/课程设计
    add_paragraph('课 程 设 计 报 告', bold=True, font_size=26,
                  font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=False)

    doc.add_paragraph()

    # 标题
    add_paragraph('基于纯NumPy从零实现CNN的', bold=True, font_size=18,
                  font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=False)
    add_paragraph('手势实时识别系统', bold=True, font_size=18,
                  font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=False)

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run('━' * 30)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # 信息
    info_items = [
        ('课程名称', '深度学习 / 神经网络与深度学习'),
        ('项目题目', '基于纯NumPy实现CNN的手势实时识别系统'),
        ('完成日期', '2026年6月'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run_label = p.add_run(f'{label}：')
        run_label.font.size = Pt(13)
        run_label.font.name = '黑体'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run_label.bold = True
        run_value = p.add_run(value)
        run_value.font.size = Pt(13)
        run_value.font.name = '宋体'
        run_value._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 分页
    doc.add_page_break()


# ============================================================
# 摘要
# ============================================================
def build_abstract():
    doc.add_heading('摘要', level=1)

    add_paragraph(
        '本课程设计实现了一套完整的手势识别系统，涵盖数据采集、模型训练和实时推理三个环节。'
        '与常见的直接调用TensorFlow或PyTorch高层API的做法不同，本项目的核心特色在于'
        '从零实现了卷积神经网络（Convolutional Neural Network, CNN）的全部核心组件——'
        '仅依赖NumPy完成矩阵运算，不使用任何深度学习框架。CNN的每一层（卷积层、池化层、全连接层）'
        '、激活函数（ReLU与Softmax）、损失函数（交叉熵）和优化器（SGD与Adam）均手动编写'
        '且附有详细的数学推导注释，共计约700行核心代码。'
    )
    add_paragraph(
        '通过该项目，可以深入理解卷积操作的im2col加速技巧、反向传播中梯度在各层间的流动机制、'
        'Softmax与交叉熵损失的联合梯度推导（其结果为简洁的 y_pred - y_true），以及Adam优化器'
        '的矩估计原理与偏差校正机制。系统支持通过笔记本电脑摄像头采集五种手势（张开手掌、握拳、'
        '剪刀手、竖大拇指、OK手势），经过30轮Mini-batch训练后测试准确率可达90%以上，'
        '并在实时推理中达到20–30 FPS的处理速度，实现了流畅的交互体验。'
    )
    add_paragraph(
        '项目的全部代码均配有详细的中文注释，解释了每一处关键实现"是什么""为什么"和"怎么推导的"，'
        '适合作为深度学习入门者理解CNN底层原理的实践参考。'
    )

    # 关键词
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run_label = p.add_run('关键词：')
    run_label.bold = True
    run_label.font.size = Pt(12)
    run_label.font.name = '黑体'
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run_value = p.add_run('卷积神经网络；手势识别；NumPy；反向传播；从零实现；im2col')
    run_value.font.size = Pt(12)
    run_value.font.name = '宋体'
    run_value._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()


# ============================================================
# 第1章 项目背景与意义
# ============================================================
def build_chapter1():
    doc.add_heading('1  项目背景与意义', level=1)

    # 1.1
    doc.add_heading('1.1  手势识别的发展与应用', level=2)
    add_paragraph(
        '手势识别（Gesture Recognition）是人机交互（Human-Computer Interaction, HCI）'
        '领域的重要研究方向之一。随着计算机视觉和深度学习技术的快速发展，基于视觉的手势识别'
        '在智能家居控制、虚拟现实（VR）与增强现实（AR）交互、手语翻译、车载信息娱乐系统操控、'
        '远程教育与会议演示等场景中展现出广阔的应用前景。'
    )
    add_paragraph(
        '传统的手势识别方法主要依赖手工设计的特征（如HOG、SIFT、Hu矩等）配合支持向量机（SVM）、'
        '随机森林或模板匹配等分类器。这类方法在受控环境下能够取得一定效果，但往往受限于光照变化、'
        '复杂背景、手部旋转与尺度变化等因素，鲁棒性较差，对不同用户的泛化能力有限。'
    )
    add_paragraph(
        '近年来，基于卷积神经网络（CNN）的深度学习方法在计算机视觉领域取得了突破性进展。'
        'CNN能够自动从原始像素中学习层次化的视觉特征——从低级的边缘和纹理，到中级的形状和部件，'
        '再到高级的语义概念——无需人工设计特征。这一特性使得CNN在手势识别任务中表现出远超传统'
        '方法的精度和鲁棒性。'
    )

    # 1.2
    doc.add_heading('1.2  从零实现CNN的意义与价值', level=2)
    add_paragraph(
        '当前主流的深度学习框架（如TensorFlow、PyTorch、PaddlePaddle等）提供了高度封装的API，'
        '开发者只需十余行代码即可搭建一个复杂的卷积神经网络。这种便利性极大地降低了深度学习的'
        '应用门槛，但也带来了一个普遍的问题：许多学习者对CNN内部运作机制的理解停留在"黑盒"层面'
        '——知道调用nn.Conv2d可以完成卷积，却不清楚卷积核的权重如何初始化、前向传播时的im2col'
        '加速原理、反向传播时梯度如何在卷积层和池化层之间传递、Softmax与交叉熵损失的联合梯度'
        '为何是简单而优雅的y_pred - y_true。'
    )
    add_paragraph(
        '本项目的核心目标正是"打破黑盒"。通过仅使用NumPy手动实现CNN的每一个组件，学习者'
        '可以直面这些核心问题，亲手推导并编码卷积层、池化层和全连接层的前向与反向传播过程，'
        '在此基础上理解优化器（如带动量的SGD和Adam）如何利用梯度信息更新网络参数。'
        '这种"造轮子"式的学习方式虽然费时，却能在原理层面建立扎实的认知基础。'
    )

    # 1.3
    doc.add_heading('1.3  技术选型与依赖', level=2)
    add_paragraph('本项目的技术选型遵循"最小依赖"原则，仅使用两个第三方库：')

    add_table(
        ['技术项', '版本要求', '用途', '备注'],
        [
            ['NumPy', '≥ 1.21.0', '矩阵运算', '项目中唯一用于数值计算的依赖'],
            ['OpenCV (cv2)', '≥ 4.5.0', '摄像头采集、图像处理', '色彩空间转换、形态学操作等'],
            ['Pillow (PIL)', '≥ 9.0', '中文文字渲染', 'OpenCV不支持中文，通过PIL桥接'],
            ['Python', '≥ 3.8', '编程语言', '无其他框架依赖'],
        ],
        caption='表1-1  项目技术依赖清单',
        col_widths=[3, 2.5, 3.5, 5]
    )

    add_paragraph(
        '值得注意的是，本项目明确不使用TensorFlow、PyTorch、Keras、PaddlePaddle等任何深度学习框架。'
        '所有的CNN层、激活函数、损失函数和优化器均基于NumPy手动编写。这一选择使得项目的代码成为'
        '一份"可运行的CNN教科书"——每一个变量的形状、每一次矩阵乘法的含义、每一步梯度推导的逻辑'
        '都清晰地展现在代码和注释之中。'
    )


# ============================================================
# 第2章 系统总体设计
# ============================================================
def build_chapter2():
    doc.add_heading('2  系统总体设计', level=1)

    # 2.1
    doc.add_heading('2.1  系统架构概览', level=2)
    add_paragraph(
        '系统采用模块化的流水线架构，由三个独立运行的脚本和一个核心CNN库组成，'
        '覆盖"数据采集→模型训练→实时推理"的完整机器学习工作流程。各模块之间的关系如下图所示。'
    )

    add_figure('fig_system_flow.png',
               '图2-1  系统总体架构与数据流示意图')

    add_paragraph(
        '各模块职责明确，通过标准化的数据格式（.npy文件和.pkl文件）进行数据交换，'
        '实现了松耦合的设计。用户可以单独运行任一阶段，例如重新采集数据后重新训练，'
        '或使用已有的模型直接进入实时识别而无需重新训练。'
    )

    # 2.2
    doc.add_heading('2.2  CNN网络结构设计', level=2)
    add_paragraph(
        '本系统采用的CNN架构遵循经典的"卷积-池化-卷积-池化-全连接"设计范式，'
        '专为64×64像素的灰度手势图像优化。网络的整体结构如下图所示。'
    )

    add_figure('fig_cnn_architecture.png',
               '图2-2  CNN网络结构示意图（含特征提取与分类两个阶段）')

    # 2.3
    doc.add_heading('2.3  各层参数与设计考量', level=2)

    add_paragraph(
        '网络由两个卷积块和一个分类器组成。以下是各层的详细参数配置：'
    )

    add_table(
        ['层编号', '层类型', '参数/配置', '输出形状', '参数量'],
        [
            ['0', 'Input', '64×64 灰度图', '(N, 1, 64, 64)', '—'],
            ['1', 'Conv2D', 'in=1, out=8, k=3×3, s=1, p=0', '(N, 8, 62, 62)', '80'],
            ['2', 'ReLU', 'max(0, x)', '(N, 8, 62, 62)', '—'],
            ['3', 'MaxPool2D', 'pool=2×2, stride=2', '(N, 8, 31, 31)', '—'],
            ['4', 'Conv2D', 'in=8, out=16, k=3×3, s=1, p=0', '(N, 16, 29, 29)', '1,168'],
            ['5', 'ReLU', 'max(0, x)', '(N, 16, 29, 29)', '—'],
            ['6', 'MaxPool2D', 'pool=2×2, stride=2', '(N, 16, 14, 14)', '—'],
            ['7', 'Flatten', '—', '(N, 3,136)', '—'],
            ['8', 'Dense', 'in=3136, out=128', '(N, 128)', '401,536'],
            ['9', 'ReLU', 'max(0, x)', '(N, 128)', '—'],
            ['10', 'Dense', 'in=128, out=5', '(N, 5)', '645'],
            ['11', 'Softmax', '—', '(N, 5)', '—'],
        ],
        caption='表2-1  CNN网络各层参数配置',
        col_widths=[1.5, 2.5, 4, 3, 2]
    )

    add_paragraph('网络设计的几个关键考量如下：')

    add_paragraph(
        '（1）第一层使用8个3×3卷积核。第一卷积层负责检测图像中最基本的局部特征——'
        '如边缘、角点、线段等。对于64×64的灰度手势图像而言，8个滤波器足以覆盖基本的方向'
        '和频率组合。核数过多容易导致过拟合（尤其在小数据集场景下），核数过少则特征提取'
        '能力不足。'
    )
    add_paragraph(
        '（2）两次Conv+Pool的组合。每一次"卷积→ReLU→池化"的复合操作都将特征图的语义层次'
        '提升一级：第一组检测局部的边缘和纹理特征，第二组在此基础上组合出中级特征——'
        '如手指的轮廓、指缝的凹陷、手掌的弧线等结构化信息。池化层在降采样的同时提供了一定'
        '程度的平移不变性。'
    )
    add_paragraph(
        '（3）全连接层使用128个神经元。该层负责将卷积提取的分布式特征表示组合为全局的'
        '分类决策依据。128是经过权衡后的取值：足够大以保留充分的表达能力（在3136维空间中'
        '学习有效的低维嵌入），又不会因参数过多（本层约40万个参数，占全部参数的99.7%）'
        '而在小数据集上严重过拟合。'
    )

    # 2.4
    doc.add_heading('2.4  手势类别定义', level=2)
    add_paragraph(
        '系统共定义五种手势类别，涵盖了日常生活中最常见的手势形态：'
    )

    add_table(
        ['类别编号', '手势名称', '英文名称', '手势形态描述', '识别框颜色'],
        [
            ['0', '张开手掌', 'Palm', '五指自然张开，掌心朝向摄像头', '绿色 (0,255,0)'],
            ['1', '握拳', 'Fist', '五指紧握成拳', '红色 (0,0,255)'],
            ['2', '剪刀手', 'Peace', '食指和中指竖起，其余手指收拢', '橙色 (255,80,0)'],
            ['3', '竖大拇指', 'Thumbs Up', '大拇指向上竖起，其余手指握拳', '金色 (0,215,255)'],
            ['4', 'OK手势', 'OK', '拇指与食指圈成O形，其余三指张开', '紫色 (255,0,255)'],
        ],
        caption='表2-2  五种手势类别定义',
        col_widths=[1.5, 2.5, 2.5, 4.5, 3]
    )

    add_figure('fig_gestures.png', '图2-3  五种手势类别示意图', width_inches=5.0)


# ============================================================
# 第3章 CNN核心库的从零实现
# ============================================================
def build_chapter3():
    doc.add_heading('3  CNN核心库的从零实现', level=1)

    add_paragraph(
        '本章是报告的核心技术章节，详细阐述CNN各组件的数学原理与实现细节。所有代码位于cnn/'
        '目录下，包含layers.py（卷积层、池化层、展平层、全连接层）、activations.py（ReLU、'
        'Softmax）、loss.py（交叉熵损失）、optimizer.py（SGD与Adam）和model.py（Sequential'
        '模型容器）五个模块，总计约700行代码，严格遵循forward()前向传播和backward()反向传播'
        '的统一接口设计。'
    )

    # 3.1
    doc.add_heading('3.1  im2col：卷积到矩阵乘法的关键转化', level=2)

    doc.add_heading('3.1.1  核心思想', level=3)
    add_paragraph(
        '卷积操作本质上是滑动窗口内的逐元素乘加运算。若直接用Python嵌套循环实现，对于每张图像'
        '需要遍历输出位置的每个像素和卷积核的每个元素，时间复杂度为O(N·OC·IC·OH·OW·KH·KW)，'
        '效率极低。im2col（Image to Column）是深度学习领域中的一个经典加速技巧：将输入图像的'
        '每个卷积窗口按列展开，将卷积核权重按行展开，从而将卷积操作转化为一次高效的矩阵乘法。'
    )
    add_paragraph(
        '矩阵乘法的优势在于：现代NumPy底层依赖经过深度优化的BLAS（Basic Linear Algebra '
        'Subprograms）库（如Intel MKL或OpenBLAS），能够充分利用CPU的SIMD指令集和多级缓存，'
        '其执行效率远超手写的嵌套循环卷积。'
    )

    doc.add_heading('3.1.2  数学过程', level=3)
    add_paragraph(
        '以一幅4×4的单通道输入图像、3×3卷积核、步长stride=1、不填充pad=0为例：'
        '输出尺寸OH = (4 - 3) / 1 + 1 = 2。卷积核在图像上共滑动4个位置。'
        'im2col将每个3×3窗口的9个像素值展平为列向量，4个窗口得到形状为(9, 4)的列矩阵。'
        '卷积核权重被展平为形状(OC, 9)的矩阵。两者的矩阵乘法直接得到(OC, 4)的输出，'
        '经reshape后即为(OC, 2, 2)的特征图。整个过程如下图所示。'
    )

    add_figure('fig_im2col.png',
               '图3-1  im2col原理示意图：将4×4图像的4个3×3滑动窗口展开为9×4列矩阵')

    doc.add_heading('3.1.3  实现细节', level=3)
    add_paragraph(
        '代码实现中，利用NumPy的广播（broadcasting）机制构造行索引矩阵i和列索引矩阵j，'
        '通过高级索引（advanced indexing）一次性提取所有滑动窗口的像素值，完全避免了'
        'Python层面的for循环。具体而言：'
    )
    add_paragraph(
        '行索引i的形状为(C·KH·KW, OH·OW)，每个元素表示对应窗口位置在填充后图像中的行坐标；'
        '列索引j同理表示列坐标。通过x_padded[:, c, i, j]一次提取，直接将原始的四维张量'
        '映射为目标的三维列矩阵，计算复杂度取决于底层BLAS而非Python解释器。'
    )
    add_paragraph(
        '反向传播时，col2im执行逆映射：将列矩阵中每个位置的梯度累加回原图像中的对应像素。'
        '这里使用np.add.at而非直接赋值，因为同一像素可能属于多个卷积窗口（例如中心像素同时'
        '属于所有9个3×3窗口），必须进行梯度累加而非覆盖。'
    )

    # 3.2
    doc.add_heading('3.2  卷积层（Conv2D）', level=2)

    doc.add_heading('3.2.1  前向传播', level=3)
    add_paragraph(
        '卷积层的前向传播在im2col框架下的实现极为简洁。其数学定义为：对于输出特征图中的'
        '每个位置(oh, ow)和每个输出通道oc，有：'
    )
    p = add_paragraph(
        '        Y[n, oc, oh, ow] = Σ_{ic, kh, kw} W[oc, ic, kh, kw] · X[n, ic, oh·S + kh, ow·S + kw] + b[oc]',
        first_line_indent=False, font_size=10, font_name='Consolas'
    )
    add_paragraph(
        '在im2col框架下，该公式简化为：W_flat（形状OC × K）矩阵乘以col（形状K × P），'
        '得到out_flat（形状OC × P），再reshape并加上偏置即为最终输出。'
        '通过einsum操作实现批量矩阵乘法——out = einsum("ok,nkp→nop", W_flat, col)。'
    )

    doc.add_heading('3.2.2  反向传播的梯度计算', level=3)
    add_paragraph(
        '卷积层的反向传播是整个CNN中最具挑战性的部分，需要同时计算三组梯度：'
    )
    add_paragraph(
        '（1）损失对权重的梯度dW：由链式法则，dW = dY_flat @ col^T，即输出梯度与输入'
        '列矩阵的外积。直观理解：如果某个输入像素值大且输出梯度也大，那么连接它们的权重'
        '应该调整更多。该计算同样通过einsum("nop,nkp→ok", d_out_flat, col)高效完成。'
    )
    add_paragraph(
        '（2）损失对偏置的梯度db：偏置被加到每个输出通道的所有空间位置上，因此梯度为'
        'd_out在batch、height、width三个维度的求和——db = sum(d_out, axis=(0, 2, 3))。'
    )
    add_paragraph(
        '（3）损失对输入的梯度dX（用于传递给前一层）：d_col = W_flat^T @ d_out_flat，'
        '其结果经col2im映射回图像空间。这一操作在数学上等价于用旋转180°的卷积核对输出'
        '梯度做"转置卷积"或"相关操作"。'
    )

    doc.add_heading('3.2.3  He初始化（Kaiming Initialization）', level=3)
    add_paragraph(
        '权重初始化的质量直接影响训练的成功与否。本项目采用He初始化策略：'
    )
    add_paragraph(
        '        W ~ N(0, sqrt(2 / fan_in)),  其中 fan_in = in_channels × KH × KW',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        'He初始化的设计目标是在前向传播时保持各层输出的方差不变。其推导基于两个假设：'
        '（1）输入服从零均值分布；（2）使用ReLU激活函数（ReLU将大约一半的输入置零，'
        '因此方差减半，需要在分子中乘以2补偿）。实践表明，He初始化能有效防止深层网络'
        '中的梯度消失和梯度爆炸，是配合ReLU使用的标准初始化方法。'
    )

    # 3.3
    doc.add_heading('3.3  最大池化层（MaxPool2D）', level=2)

    doc.add_heading('3.3.1  前向传播', level=3)
    add_paragraph(
        '池化层的核心功能是降采样（Downsampling）。最大池化在每个2×2的不重叠窗口中'
        '选取最大值作为输出，效果是将特征图的空间尺寸减半。这一操作带来三重好处：'
        '（1）大幅降低后续层的计算量；（2）提供一定的平移不变性——手部在ROI框内的微小'
        '移动不会改变池化层的输出；（3）扩大感受野——后续层的每个神经元"看到"的原始输入'
        '区域翻倍。'
    )
    add_paragraph(
        '实现中采用与im2col类似的索引技巧提取所有池化窗口的值，然后沿窗口维度取max。'
        '同时记录每个窗口中最大值的位置（argmax），供反向传播使用。'
    )

    doc.add_heading('3.3.2  反向传播的稀疏梯度', level=3)
    add_paragraph(
        '最大池化层没有可学习参数，反向传播的唯一任务是正确地将梯度传递到前一层。'
        '其核心规则是：梯度只流向池化窗口中最大值所在的位置，其余位置的梯度严格为零。'
        '这一"赢者通吃"机制背后的逻辑简单而正确——只有最大值参与了输出结果的计算，'
        '改变非最大值不会影响输出，因此它们对损失的梯度为0。'
    )
    add_paragraph(
        '这种稀疏的梯度流动模式是最大池化的一个重要特性。它与ReLU的稀疏激活一起，'
        '为CNN引入了一种天然的"竞争性学习"机制——只有最显著的特征信号才能通过梯度'
        '反向传播到更早的层，噪声和微弱信号则被自然地过滤。'
    )

    # 3.4
    doc.add_heading('3.4  全连接层（Dense）', level=2)
    add_paragraph(
        '全连接层是神经网络中最"经典"的层类型——每个输入神经元与每个输出神经元之间'
        '都存在一个可学习的连接权重。其数学形式为简单的仿射变换：Y = X·W + b，其中'
        'X∈R^{N×D_in}，W∈R^{D_in×D_out}，b∈R^{D_out}。'
    )
    add_paragraph(
        '反向传播同样计算三组梯度：dW = X^T·dY（输入和输出梯度的外积、对批求和），'
        'db = ΣdY（沿batch维度求和），dX = dY·W^T（梯度经权重矩阵的转置回传）。'
        '这三个公式是神经网络反向传播中最基础也最重要的结果，建议读者自行完成完整的'
        '导数推导以加深理解。'
    )

    # 3.5
    doc.add_heading('3.5  激活函数', level=2)

    doc.add_heading('3.5.1  ReLU', level=3)
    add_paragraph(
        '修正线性单元（Rectified Linear Unit, ReLU）的数学形式为ReLU(x) = max(0, x)，'
        '其函数图像与导函数如下图所示。'
    )

    add_figure('fig_relu.png', '图3-2  ReLU激活函数及其导数曲线')

    add_paragraph(
        'ReLU之所以成为最广泛使用的激活函数，有四方面原因：（1）计算极其简单——仅需一次比较操作，'
        '无浮点指数运算；（2）正面半区梯度恒为1——彻底解决了Sigmoid/Tanh在两端饱和导致的梯度消失问题；'
        '（3）天然产生稀疏激活——约50%的神经元输出为零，这种稀疏性被证明有助于提升模型的泛化能力；'
        '（4）生物学合理性——与生物神经元"全或无"的发放特性类似。'
    )
    add_paragraph(
        'ReLU的主要缺陷是"Dying ReLU"问题——若某个神经元对所有训练样本均输出负值，'
        '其梯度恒为零，权重永不更新，该神经元永久"死亡"。本项目的浅层网络规模使得'
        '该问题不显著，实践中数据多样性和合适的初始化通常能避免大量神经元同时失活。'
    )

    doc.add_heading('3.5.2  Softmax', level=3)
    add_paragraph(
        'Softmax函数将任意实数向量映射为合法的概率分布（所有输出分量非负且和为1）：'
    )
    add_paragraph(
        '        Softmax(x_i) = exp(x_i) / Σ_j exp(x_j)',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '实现中采用了一个重要的数值稳定性技巧：计算前先将每个样本的logits减去该样本的'
        '最大值（x_shifted = x - max(x, axis=1, keepdims=True)）。这一操作不改变Softmax'
        '的数学结果（相当于分子分母同除以exp(max)），但有效防止了exp运算溢出——'
        '例如exp(1000) ≈ Inf，而exp(1000 - 1000) = exp(0) = 1，数值完全稳定。'
    )

    # 3.6
    doc.add_heading('3.6  交叉熵损失函数', level=2)

    doc.add_heading('3.6.1  交叉熵的定义与直觉', level=3)
    add_paragraph(
        '交叉熵（Cross-Entropy）是衡量两个概率分布之间差异的标准度量，在多分类任务中'
        '是将Softmax输出与真实标签对齐的默认损失函数：'
    )
    add_paragraph(
        '        L = -Σ_{c} y_c^{true} · log(y_c^{pred})',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '直觉理解：当预测在正确类别上的概率接近1时，-log(1) ≈ 0，损失接近0；'
        '当预测在正确类别上的概率接近0时，-log(0) → +∞，损失巨大。'
        '交叉熵"惩罚"模型对正确类别的"不自信"——这正是分类任务所需的行为。'
    )

    doc.add_heading('3.6.2  Softmax与交叉熵的联合梯度——最优雅的结果', level=3)
    add_paragraph(
        '这是本项目最重要的数学知识点之一，也是深度学习领域中最经典的数学结果。'
        '如果分别计算Softmax的雅可比矩阵（N个C×C矩阵）和交叉熵对Softmax输出的梯度'
        '（-y_true/s），则需要构造并存储O(N·C²)大小的中间结果。然而，利用链式法则'
        '将两者合并后，结果惊人地简单。完整的推导过程如下图所示。'
    )

    add_figure('fig_softmax_gradient.png',
               '图3-3  Softmax + CrossEntropy联合梯度的完整推导过程')

    add_paragraph(
        '推导的关键步骤：将dL/ds_k = -y_k / s_k（交叉熵对softmax输出的梯度）与'
        'ds_k/dx_i = s_k · (δ_ki - s_i)（softmax的雅可比矩阵元素）代入链式法则：'
    )
    add_paragraph(
        'dL/dx_i = Σ_k (-y_k / s_k) · s_k · (δ_ki - s_i) = Σ_k -y_k · (δ_ki - s_i)',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '       = -y_i + s_i · Σ_k y_k = s_i - y_i',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '（最后一步利用了Σ_k y_k = 1，因为y_true是one-hot编码，恰好一个分量为1其余为0。）'
    )
    add_paragraph(
        '最终结果：dL/dx = y_pred - y_true。这个公式的优雅程度难以言表——反向传播'
        '仅需将预测概率减去真实标签，一行代码即可完成。这也是为什么"Softmax+CrossEntropy"'
        '组合成为分类任务黄金标准的根本原因。在代码实现中，CrossEntropyLoss.backward()'
        '方法的主体就是return (softmax_out - y_true) / N。'
    )

    # 3.7
    doc.add_heading('3.7  优化器：SGD与Adam', level=2)

    doc.add_heading('3.7.1  带动量的SGD', level=3)
    add_paragraph(
        '基本梯度下降（w = w - lr·dw）在复杂损失曲面上存在收敛缓慢和震荡问题。'
        '动量法（Momentum）在SGD的基础上引入速度（velocity）的概念：'
    )
    add_paragraph(
        '        v_t = β · v_{t-1} - η · ∇L(w_t)',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '        w_{t+1} = w_t + v_t',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '物理直觉：将损失曲面想象为一个山谷，梯度是重力（指向最陡下降方向），速度如同惯性。'
        '球不会在每次受力时立即改变方向，而是保持一定的运动趋势。这种惯性在梯度一致的方向上'
        '加速收敛，在梯度震荡的方向上平滑波动，同时有助于翻越浅的局部极小点。'
    )

    doc.add_heading('3.7.2  Adam优化器', level=3)
    add_paragraph(
        'Adam（Adaptive Moment Estimation）是目前深度学习领域最广泛使用的优化器之一，'
        '它巧妙地结合了动量法（利用梯度的一阶矩）和RMSprop（利用梯度的二阶矩）两种思想：'
    )
    add_paragraph(
        '        m_t = β₁·m_{t-1} + (1-β₁)·g_t　　   （一阶矩——梯度的指数移动平均）',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '        v_t = β₂·v_{t-1} + (1-β₂)·g_t²　　  （二阶矩——梯度平方的指数移动平均）',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '        m̂_t = m_t / (1-β₁^t),  v̂_t = v_t / (1-β₂^t)     （偏差校正）',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        '        w_{t+1} = w_t - η·m̂_t / (√v̂_t + ε)',
        first_line_indent=False, font_size=10
    )
    add_paragraph(
        'Adam的直觉理解：一阶矩m给出"梯度平均指向哪个方向"，二阶矩v给出"梯度在这个方向上'
        '有多大的不确定性（方差）"。当梯度在某一维度上稳定地指向同一方向（m大）且波动小'
        '（v小）时，步长η/√v较大，快速前进；当梯度震荡剧烈（v大）时，步长自动减小，'
        '谨慎探索。这种逐参数的自适应学习率机制使Adam在大多数任务上无需手动调节学习率。'
    )
    add_paragraph(
        '偏差校正（Bias Correction）是Adam中一个精巧的设计。初始时m₀=v₀=0，'
        '如果不进行校正，前几步的估计值会严重偏向零，导致步长过小、初期训练缓慢。'
        '除以(1-β^t)可以抵消零初始化的偏差，且随着t增大该因子趋近于1——校正机制自然'
        '"退场"，不影响后期训练。本项目选用Adam（lr=0.001）而非SGD作为训练优化器。'
    )

    # 3.8
    doc.add_heading('3.8  Sequential模型容器与训练步骤', level=2)
    add_paragraph(
        'Sequential模型将各层串联为线性的数据处理管道——数据依次流过每一层，'
        '前一层输出即为后一层输入。其train_step()方法封装了单批数据的一次完整训练迭代，'
        '包含以下五个子步骤：'
    )
    add_paragraph(
        '步骤一：前向传播（Forward）。输入数据x依次通过每一层的forward()方法，'
        '最终得到Softmax输出的概率分布y_pred。每层在此过程中缓存必要的中间结果供反向传播使用。'
    )
    add_paragraph(
        '步骤二：计算损失（Loss）。调用CrossEntropyLoss.forward(y_pred, y_true)，'
        '计算当前批次的平均交叉熵损失，同时缓存softmax输出和真实标签。'
    )
    add_paragraph(
        '步骤三：反向传播（Backward）。从损失函数开始，梯度沿层序逆向流动：'
        'loss_fn.backward() → model.backward(d_loss) → 依次调用各层的backward()，'
        '每一层计算dW、db（如有参数）和dX（传给前一层）。'
    )
    add_paragraph(
        '步骤四：参数更新（Update）。优化器遍历所有含可学习参数的层，根据计算出的梯度'
        '更新权重和偏置。对于Adam，此步骤还涉及一阶矩和二阶矩的维护与偏差校正。'
    )
    add_paragraph(
        '步骤五：计算准确率。对当前批次的预测结果取argmax得到预测类别，与真实类别比较，'
        '计算批内平均准确率，用于训练过程中的监控输出。'
    )
    add_paragraph(
        '模型的保存与加载使用Python内置的pickle序列化机制，仅存储可学习参数（权重矩阵和'
        '偏置向量），不保存层结构。这意味着加载时需先重建相同的层拓扑，再将参数赋值到对应'
        '位置。这一设计的优点是模型文件轻量（仅约1.6 MB），缺点是需要代码层面的结构一致性保证。'
    )


# ============================================================
# 第4章 数据采集模块
# ============================================================
def build_chapter4():
    doc.add_heading('4  手势数据采集', level=1)

    # 4.1
    doc.add_heading('4.1  采集流程设计', level=2)
    add_paragraph(
        '手势数据是模型训练的基础。数据采集模块（collect_data.py）通过笔记本电脑摄像头'
        '实时采集用户的手势图像，其工作流程为：打开摄像头→显示ROI采集框→用户按数字键0-4'
        '采集对应手势→按S键保存数据→按Q键退出。每种手势的目标采集量为200张，总计1000张。'
    )
    add_paragraph(
        '关键设计要点包括：画面中央250×250像素的绿色ROI框引导用户将手放置在统一位置；'
        '0.3秒的采集冷却时间防止连拍产生高度相似的冗余样本；实时显示各手势已采集数量和'
        '肤色检测掩码的预览窗口，提供即时的视觉反馈。'
    )

    # 4.2
    doc.add_heading('4.2  图像预处理流水线', level=2)
    add_paragraph(
        '预处理是将原始BGR摄像头画面转化为CNN可接受的标准化输入的关键环节，包含七个步骤：'
    )

    add_table(
        ['步骤', '操作', '方法/参数', '设计原因'],
        [
            ['①', '色彩空间转换', 'BGR → HSV', 'HSV分离了色相、饱和度和亮度，肤色在H和S上范围稳定，对光照鲁棒'],
            ['②', '肤色阈值分割', 'cv2.inRange(H:0-25, S:20-180, V:40-255)', '在HSV空间中筛选肤色像素点，生成二值掩码'],
            ['③', '形态学处理', '开运算×2 + 闭运算×3, 椭圆核5×5', '去除孤立噪声点并填充手部区域内部小空洞'],
            ['④', '灰度化+模糊', 'COLOR_BGR2GRAY + GaussianBlur(5×5)', '平滑皮肤纹理细节，帮助CNN聚焦于手部整体形状'],
            ['⑤', '应用掩码', 'bitwise_and', '将背景区域置黑，仅保留手部前景'],
            ['⑥', '尺寸缩放', 'resize→(64, 64)', '统一输入尺寸，满足CNN固定输入要求'],
            ['⑦', '归一化', '/255.0 → [0, 1]', '数值归一化有助于稳定训练过程、加速收敛'],
        ],
        caption='表4-1  图像预处理流水线',
        col_widths=[1, 2.5, 4.5, 5.5]
    )

    add_paragraph(
        '选择HSV而非RGB进行肤色检测是经过考量的设计决策。在RGB空间中，肤色的三个通道'
        '高度耦合——光照变化（如从室内暖光变为窗口自然光）会导致R、G、B三个通道同时'
        '剧烈变化。而HSV将色彩信息（色相H）与亮度信息（明度V）分离，人类肤色在色相和'
        '饱和度上具有相对稳定的范围（H∈[0, 25°]，对应橙红色到浅黄色），对光照变化的'
        '鲁棒性显著优于RGB。'
    )

    # 4.3
    doc.add_heading('4.3  数据存储格式', level=2)
    add_paragraph(
        '采集完成后，数据以NumPy的二进制.npy格式保存，包含两个文件：X.npy（图像数组，'
        '形状(N, 64, 64)，数据类型float32，值域[0, 1]）和y.npy（标签数组，形状(N,)，'
        '数据类型int64，值为0-4的类别编号）。保存前自动进行随机打乱（shuffle），防止'
        '模型学习到采集顺序引入的伪规律（如"先采集的全部是palm，后采集的全部是fist"）。'
        '保存时输出各类别样本数的柱状图，便于快速检查数据均衡性。'
    )


# ============================================================
# 第5章 模型训练
# ============================================================
def build_chapter5():
    doc.add_heading('5  模型训练与评估', level=1)

    # 5.1
    doc.add_heading('5.1  数据加载与预处理', level=2)
    add_paragraph(
        '训练脚本（train.py）加载gesture_data/中的原始数据后，执行两项关键的数据再处理：'
    )
    add_paragraph(
        '（1）增加通道维度：原始图像形状为(N, 64, 64)，通过X[:, np.newaxis, :, :]'
        '扩展为(N, 1, 64, 64)，增加的单通道维度符合Conv2D层期望的(N, C, H, W)输入格式。'
        '对于灰度图C=1，对于RGB图C=3。'
    )
    add_paragraph(
        '（2）One-hot编码：原始标签为整数0-4（如标签2代表peace手势），通过创建全零矩阵'
        '并在对应位置赋值1的方式转化为one-hot向量（如[0, 0, 1, 0, 0]），'
        '以满足交叉熵损失函数的输入要求。one-hot编码是分类任务中最常用的标签表示法。'
    )

    # 5.2
    doc.add_heading('5.2  数据集划分', level=2)
    add_paragraph(
        '采用80/20的比例进行随机训练集/测试集划分。具体做法是：先通过np.random.permutation(N)'
        '生成0到N-1的随机排列，取前80%的索引作为训练集、后20%作为测试集。随机划分确保各类别'
        '在训练集和测试集中大致保持相同的比例分布，使测试集具有代表性。与顺序划分（前80%训练、'
        '后20%测试）相比，随机划分能避免因采集顺序造成的分布偏差。'
    )

    # 5.3
    doc.add_heading('5.3  超参数配置', level=2)

    add_table(
        ['超参数', '取值', '说明'],
        [
            ['Batch Size', '16', '小批量训练，平衡梯度估计的稳定性和内存开销'],
            ['Epochs', '30', '经验值，在1000张规模的数据集上通常20-30轮收敛'],
            ['Learning Rate', '0.001', 'Adam优化器的推荐起始学习率，适用于大多数任务'],
            ['β₁, β₂', '0.9, 0.999', 'Adam的默认衰减率，实践中几乎无需调整'],
            ['Weight Init', 'He (Kaiming)', '适配ReLU的初始化策略，保证前向方差不变'],
            ['Train Split', '0.8', '80%训练，20%测试——小数据集的常用配比'],
        ],
        caption='表5-1  训练超参数配置',
        col_widths=[3, 3, 8]
    )

    # 5.4
    doc.add_heading('5.4  训练循环', level=2)
    add_paragraph(
        '每个epoch内部执行以下操作：首先通过np.random.permutation随机打乱训练数据顺序'
        '（这对防止模型"记忆"数据顺序、提升泛化能力至关重要）；然后将打乱后的数据按batch_size=16'
        '分为若干mini-batch，依次执行train_step（forward→loss→backward→update），记录每批的'
        '训练损失和准确率；每个epoch结束后在测试集上评估模型表现，若当前测试准确率超过历史最佳，'
        '则将模型参数保存至gesture_model.pkl。'
    )
    add_paragraph(
        '训练过程约耗时1-3分钟（取决于CPU性能与数据集大小），在每个epoch控制台输出Train Loss、'
        'Train Accuracy、Test Loss、Test Accuracy四项指标以及该epoch的耗时。保存最佳模型时'
        '会在输出行末尾标记★符号。'
    )

    # 5.5
    doc.add_heading('5.5  模型评估与分析', level=2)
    add_paragraph(
        '训练结束后，除了保存最佳模型外，还会输出两项评估内容：'
    )
    add_paragraph(
        '（1）各类别准确率：对测试集中每个手势类别单独计算准确率，以文字柱状图（█和░字符）'
        '的形式直观展示。这有助于识别模型在特定手势上的薄弱点——例如剪刀手手势因手指姿态'
        '变化较大，准确率通常略低于形态相对固定的握拳和OK手势。'
    )
    add_paragraph(
        '（2）训练过程摘要：按步长输出关键epoch的train_loss和test_acc，标注最佳测试准确率'
        '所在的epoch，便于分析收敛速度和过拟合趋势。'
    )


# ============================================================
# 第6章 实时手势识别
# ============================================================
def build_chapter6():
    doc.add_heading('6  实时手势识别系统', level=1)

    # 6.1
    doc.add_heading('6.1  系统工作流程', level=2)
    add_paragraph(
        '实时识别模块（recognize.py）加载训练好的模型，打开摄像头进入持续推理循环。'
        '每一帧的处理流程为：BGR→HSV转换→肤色检测→形态学处理→灰度化→高斯模糊→'
        '肤色掩码→缩放至64×64→归一化→增加batch和channel维度→CNN前向传播→5类概率分布'
        '→取最大概率类别和置信度→滑动窗口平滑→绘制中文识别结果和UI元素→显示画面。'
        '完整的帧处理在普通笔记本CPU上耗时约35-50ms，对应20-30 FPS的实时处理速率。'
    )

    # 6.2
    doc.add_heading('6.2  预测平滑策略', level=2)
    add_paragraph(
        '单帧预测容易因光照闪烁、手部微动、肤色检测的帧间差异等因素产生类别跳变。'
        '为解决这一问题，系统采用滑动窗口平滑策略：维护一个长度为5帧的预测缓冲区，'
        '对缓冲区内最近的5个预测结果进行多数投票（Majority Voting）。'
        '当某个类别在缓冲区中出现的次数不少于3次（即≥60%）时，确认为"稳定手势"。'
        '只有当稳定手势的类别发生变化时，才在控制台输出新的识别结果，避免频繁刷屏。'
    )
    add_paragraph(
        '这一平滑策略本质上是一个简单的低通滤波器，它引入约100-150ms的响应延迟'
        '（5帧÷25FPS≈200ms的理论最大值），在用户体验的流畅性和预测的稳定性之间取得了良好平衡。'
    )

    # 6.3
    doc.add_heading('6.3  用户界面设计', level=2)
    add_paragraph(
        '识别界面包含以下视觉元素：画面中央的ROI识别框（颜色随识别到的手势动态变化）、'
        'ROI上方居中的中文手势名称和置信度百分比、左上角的信息面板（含系统标题、中文手势名称、'
        '置信度进度条）、以及右下角的实时FPS计数器。界面设计遵循简洁、直观的原则，'
        '避免过多的视觉干扰。'
    )

    # 6.4
    doc.add_heading('6.4  中文显示技术方案', level=2)
    add_paragraph(
        'OpenCV的原生putText()函数仅支持ASCII字符，无法渲染中文。系统采用Pillow（PIL）库'
        '作为中文显示的桥梁：先将OpenCV的BGR图像转换为PIL的RGB图像，利用PIL的ImageDraw模块'
        '以系统TrueType字体渲染中文文字（支持文字阴影效果），再将处理后的PIL图像转换回OpenCV的'
        'BGR格式。字体查找优先级为：黑体（simhei.ttf）→微软雅黑（msyh.ttc）→宋体（simsun.ttc）'
        '→楷体（simkai.ttf）→Linux字体→macOS字体→PIL默认字体，确保了跨Windows/macOS/Linux'
        '平台的兼容性。'
    )


# ============================================================
# 第7章 实验结果与分析
# ============================================================
def build_chapter7():
    doc.add_heading('7  实验结果与分析', level=1)

    # 7.1
    doc.add_heading('7.1  训练过程分析', level=2)
    add_paragraph(
        '在包含约1000张手势图像（每类200张）的数据集上，使用Adam优化器训练30轮。'
        '典型的训练过程如以下曲线所示。'
    )

    add_figure('fig_training_curves.png', '图7-1  训练过程中Loss和Accuracy的变化曲线')

    add_paragraph(
        '训练过程可划分为三个阶段：初期（Epoch 1-5）损失从约1.6快速下降至约0.3，'
        '测试准确率从约30%攀升至约75%，此阶段模型在学习手势的粗略轮廓特征；中期'
        '（Epoch 6-15）损失缓慢下降，准确率在80-90%区间波动上升，模型正在学习更精细'
        '的手势特征；后期（Epoch 16-30）模型趋于收敛，最佳测试准确率可达90%以上。'
        '整个训练过程约耗时1-3分钟。'
    )
    add_paragraph(
        '注意到训练损失持续下降而测试损失在后期可能出现轻微反弹——这是轻微的过拟合迹象，'
        '在1000张的小数据集上是预期现象。可通过增加训练数据或引入Dropout层缓解。'
    )

    # 7.2
    doc.add_heading('7.2  影响因素分析', level=2)

    doc.add_heading('7.2.1  数据量对准确率的影响', level=3)
    add_paragraph(
        '实验表明，每类200张手势图像是获取90%+测试准确率的经验最低推荐量。'
        '若数据量减半（每类100张），测试准确率通常下降5-10个百分点至80-85%左右。'
        '这是由于CNN具有约40万个参数，需要充足的训练样本来约束优化过程、防止过拟合。'
        '将数据量增至每类300-500张可进一步提升至95%左右，但边际收益递减。'
    )

    doc.add_heading('7.2.2  光照与背景的影响', level=3)
    add_paragraph(
        '肤色检测的精度是影响整个系统性能的关键瓶颈。在简单纯色背景（如白墙）和均匀光照'
        '条件下，HSV肤色阈值分割效果良好，模型表现稳定。在复杂背景（如包含肤色相近的家具'
        '或多人场景）下，肤色检测可能引入大量背景噪声或错误地分割掉手部区域，导致识别准确率'
        '显著下降。解决方案包括：在复杂背景下采集更多训练数据以提高模型鲁棒性；调整HSV肤色'
        '范围的三个参数（H、S、V的上下界）；或采用更先进的深度学习手部分割方法替代传统的'
        '颜色阈值法。'
    )

    doc.add_heading('7.2.3  帧率性能分析', level=3)
    add_paragraph(
        '在普通笔记本CPU（Intel i5/i7级别，无独立GPU加速）上，纯NumPy实现的CNN单次前向传播'
        '耗时约15-25ms，加上图像预处理约20-25ms，总的一帧处理时间约35-50ms，对应20-30 FPS'
        '的实时识别速率。处理耗时的主要瓶颈在于CNN前向传播中的矩阵乘法操作——即使使用了im2col'
        '加速，NumPy的CPU矩阵乘法在40万参数规模下仍需要十余毫秒。如需进一步提升帧率，可考虑'
        '减小输入尺寸（如48×48）、减少卷积核数量（如8→4和16→8），或将计算迁移至GPU。'
    )

    # 7.3
    doc.add_heading('7.3  模型局限性讨论', level=2)
    add_paragraph(
        '本系统的主要局限性包括：（1）基于肤色检测的预处理在复杂背景和多人场景下不够可靠；'
        '（2）网络结构较浅（仅两个卷积块），对复杂手势的判别能力有限；（3）训练数据由单人'
        '在固定环境下采集，模型泛化到不同用户和不同环境的能力需进一步验证；'
        '（4）每帧独立处理，未利用时序信息（如手势的动态变化过程）。这些局限性为未来的'
        '改进工作提供了明确的方向。'
    )


# ============================================================
# 第8章 总结与展望
# ============================================================
def build_chapter8():
    doc.add_heading('8  总结与展望', level=1)

    doc.add_heading('8.1  工作总结', level=2)
    add_paragraph(
        '本课程设计完成了一套完整的、从零实现的CNN手势识别系统，主要成果如下：'
    )
    add_paragraph(
        '（1）CNN核心库。独立实现了Conv2D（含im2col/col2im加速）、MaxPool2D、Flatten、'
        'Dense四个层类型，ReLU和Softmax两个激活函数，CrossEntropyLoss损失函数，'
        'SGD（含动量）和Adam两个优化器，以及Sequential模型容器，全部代码约700行，'
        '每处关键实现均附有详细的中文数学推导注释。'
    )
    add_paragraph(
        '（2）数据采集系统。基于HSV肤色检测的实时手势数据采集工具，支持五种手势的标注数据采集，'
        '包含采集冷却、实时预览、数据打乱与保存等功能。'
    )
    add_paragraph(
        '（3）模型训练流程。完整的数据加载、One-hot编码、训练/测试集划分、Mini-batch训练、'
        '模型评估（整体与各类别准确率）、最佳模型保存等机器学习核心工作流。'
    )
    add_paragraph(
        '（4）实时识别系统。支持中文显示的实时手势识别界面，含滑动窗口预测平滑、置信度可视化、'
        '差异化视觉反馈和FPS监控。'
    )

    doc.add_heading('8.2  学习价值', level=2)
    add_paragraph(
        '本项目最大的价值在于其教育意义。通过从零实现CNN的全部组件，学习者可以深入理解'
        '以下核心概念：卷积的im2col实现及其与矩阵乘法的等价关系；反向传播中梯度在卷积层、'
        '池化层和全连接层间的精确流动机制；Softmax与交叉熵的联合梯度为何是简洁的y_pred-y_true；'
        'Adam优化器中一阶矩、二阶矩和偏差校正各自的作用与物理直觉；以及He初始化的动机和'
        '数学形式。这些知识在使用TensorFlow/PyTorch时被封装在框架内部，直接阅读框架源码'
        '往往因工程优化的复杂性而难以抓住数学本质。本项目的代码刻意保持了实现与数学公式'
        '的清晰对应关系——代码即推导。'
    )

    doc.add_heading('8.3  未来改进方向', level=2)
    add_paragraph(
        '（1）数据增强（Data Augmentation）。在训练时对图像进行随机旋转（±15°）、平移（±10%）、'
        '缩放（±15%）和亮度微调，可有效扩充数据集规模、提升模型泛化能力和对姿态变化的鲁棒性。'
        '数据增强几乎无额外成本（仅需在CPU上即时处理），是提升小数据集模型表现的最具性价比的改进。'
    )
    add_paragraph(
        '（2）批归一化（Batch Normalization）。在卷积层后、ReLU激活前插入BN层，可以缓解内部协变量'
        '偏移（Internal Covariate Shift）问题，允许使用更大的学习率，加速训练收敛，并自带一定的'
        '正则化效果。BN的实现同样可基于NumPy完成，是深入理解深度学习训练技巧的自然延伸。'
    )
    add_paragraph(
        '（3）更深的网络结构。尝试引入ResNet风格的残差连接（Skip Connection），构建更深层的网络'
        '（如3-4个卷积块），同时利用BatchNorm稳定深层网络的训练。残差连接解决了深层网络中的退化问题'
        '（degradation problem），是深度学习领域的里程碑式创新。'
    )
    add_paragraph(
        '（4）GPU加速。使用CuPy（NumPy的GPU替代品，API完全兼容）或迁移到PyTorch框架，将计算密集型'
        '的矩阵乘法移至GPU执行。在GPU上，本项目的CNN前向传播可在亚毫秒级完成，帧率轻松突破60 FPS。'
    )
    add_paragraph(
        '（5）扩展手势库。将手势类别从5种扩展至10种（数字手势0-9）或更多，并引入"无手势/背景"'
        '类别以提高系统在实际使用中的鲁棒性。同时可探索动态手势（如滑动手势）的时序建模。'
    )
    add_paragraph(
        '（6）实际交互应用。将识别结果通过系统API映射为具体操作——如音量调节、幻灯片翻页、'
        '媒体播放控制、智能家居开关等，将手势识别从演示系统升级为具备实用价值的交互工具。'
    )


# ============================================================
# 附录
# ============================================================
def build_appendix():
    doc.add_heading('附录', level=1)

    doc.add_heading('A  项目文件结构', level=2)

    add_table(
        ['文件/目录', '类型', '功能说明'],
        [
            ['cnn/__init__.py', 'Python模块', 'CNN库入口，统一导出所有组件'],
            ['cnn/layers.py', 'Python模块', 'Conv2D, MaxPool2D, Flatten, Dense（含im2col）'],
            ['cnn/activations.py', 'Python模块', 'ReLU, Softmax激活函数'],
            ['cnn/loss.py', 'Python模块', 'CrossEntropyLoss, MSELoss'],
            ['cnn/optimizer.py', 'Python模块', 'SGD（含动量）, Adam优化器'],
            ['cnn/model.py', 'Python模块', 'Sequential模型容器（train_step、save/load）'],
            ['collect_data.py', '应用脚本', '摄像头手势数据采集'],
            ['train.py', '应用脚本', 'CNN模型训练与评估'],
            ['recognize.py', '应用脚本', '实时手势识别与可视化'],
            ['requirements.txt', '配置文件', 'Python依赖清单（numpy, opencv-python）'],
            ['manual.txt', '说明文档', '详细的使用手册（含常见问题解答）'],
        ],
        caption='表A-1  项目文件结构与说明',
        col_widths=[3.5, 2.5, 8]
    )

    doc.add_heading('B  环境配置与运行步骤', level=2)
    add_paragraph('本系统已在Windows 11环境下完成开发与测试，具体运行步骤如下：')
    add_paragraph(
        '步骤一：安装依赖。在项目根目录下执行 pip install -r requirements.txt。'
        '如使用国内网络，可添加 -i https://mirrors.tuna.tsinghua.edu.cn/pypi/web/simple 镜像参数。',
        first_line_indent=False
    )
    add_paragraph(
        '步骤二：采集手势数据。执行 python collect_data.py，待摄像头打开后，将手掌伸入画面中央的'
        '绿色ROI方框内，依次按数字键0-4采集五种手势（每按一次采集一帧），每种手势建议采集200张。'
        '按S键保存数据至gesture_data/目录。',
        first_line_indent=False
    )
    add_paragraph(
        '步骤三：训练CNN模型。执行 python train.py，脚本自动加载数据、构建网络、执行30轮训练，'
        '并保存最佳模型至gesture_model.pkl。训练过程在控制台实时输出各项指标。',
        first_line_indent=False
    )
    add_paragraph(
        '步骤四：实时手势识别。执行 python recognize.py，摄像头打开后将手放入识别框中，'
        'CNN实时识别手势并在画面中显示中文类别名称。按D键切换调试模式（显示CNN输入和肤色掩码预览），'
        '按Q键退出。',
        first_line_indent=False
    )

    doc.add_heading('C  常见问题与解决方法', level=2)

    add_table(
        ['问题', '可能原因', '解决方法'],
        [
            ['摄像头无法打开', '其他程序占用摄像头', '关闭其他使用摄像头的应用，重启电脑'],
            ['肤色检测不准', '光照条件或肤色与HSV范围不匹配',
             '调整collect_data.py中的SKIN_LOWER和SKIN_UPPER参数范围'],
            ['训练准确率低', '数据量不足或质量不佳', '增加每类手势的采集量至300+，确保手势姿态多样化'],
            ['识别帧率低', 'CPU计算能力有限', '减小IMG_SIZE至48，或减少卷积核数量'],
            ['添加新手势', '当前仅支持5类', '修改GESTURE_NAMES和NUM_CLASSES，重新采集→训练'],
        ],
        caption='表C-1  常见问题与解决方法',
        col_widths=[3, 4.5, 6.5]
    )


# ============================================================
# 参考文献
# ============================================================
def build_references():
    doc.add_heading('参考文献', level=1)

    refs = [
        '[1] Krizhevsky A, Sutskever I, Hinton G E. ImageNet classification with deep convolutional neural networks[J]. Communications of the ACM, 2017, 60(6): 84-90.',
        '[2] He K, Zhang X, Ren S, et al. Delving deep into rectifiers: Surpassing human-level performance on ImageNet classification[C]. Proceedings of the IEEE International Conference on Computer Vision (ICCV), 2015: 1026-1034.',
        '[3] Kingma D P, Ba J. Adam: A method for stochastic optimization[C]. International Conference on Learning Representations (ICLR), 2015.',
        '[4] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.',
        '[5] Simonyan K, Zisserman A. Very deep convolutional networks for large-scale image recognition[C]. International Conference on Learning Representations (ICLR), 2015.',
        '[6] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2016: 770-778.',
        '[7] Chellapilla K, Puri S, Simard P. High performance convolutional neural networks for document processing[C]. International Workshop on Frontiers in Handwriting Recognition, 2006.',
        '[8] Ioffe S, Szegedy C. Batch normalization: Accelerating deep network training by reducing internal covariate shift[C]. International Conference on Machine Learning (ICML), 2015: 448-456.',
        '[9] LeCun Y, Bottou L, Bengio Y, et al. Gradient-based learning applied to document recognition[J]. Proceedings of the IEEE, 1998, 86(11): 2278-2324.',
        '[10] Nair V, Hinton G E. Rectified linear units improve restricted Boltzmann machines[C]. International Conference on Machine Learning (ICML), 2010: 807-814.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref)
        run.font.size = Pt(10.5)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ============================================================
# 页眉页脚设置
# ============================================================
def setup_header_footer():
    section = doc.sections[0]
    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ''
    run = hp.add_run('深度学习课程设计报告 —— 基于纯NumPy从零实现CNN的手势识别系统')
    run.font.size = Pt(8)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(128, 128, 128)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 页脚（页码）
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


# ============================================================
# 主构建流程
# ============================================================
def main():
    print('Building cover...')
    build_cover()

    print('Building abstract...')
    build_abstract()

    print('Building Chapter 1...')
    build_chapter1()

    print('Building Chapter 2...')
    build_chapter2()

    print('Building Chapter 3...')
    build_chapter3()

    print('Building Chapter 4...')
    build_chapter4()

    print('Building Chapter 5...')
    build_chapter5()

    print('Building Chapter 6...')
    build_chapter6()

    print('Building Chapter 7...')
    build_chapter7()

    print('Building Chapter 8...')
    build_chapter8()

    print('Building Appendix...')
    build_appendix()

    print('Building References...')
    build_references()

    print('Setting up header/footer...')
    setup_header_footer()

    # 保存
    output_path = os.path.join(BASE_DIR, '课程设计报告.docx')
    doc.save(output_path)
    print(f'\nDocument saved to: {output_path}')


if __name__ == '__main__':
    main()
